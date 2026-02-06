from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Configurar margens menores para caber a tabela
for section in doc.sections:
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.page_width = Cm(29.7)   # A4 paisagem
    section.page_height = Cm(21.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

# Titulo do documento
title = doc.add_heading('Documentacao dos Servicos REST', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# Cores do cabecalho
HEADER_BG = '4472C4'
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)

def set_cell_bg(cell, color):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shading_elm = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    tcPr.append(shading_elm)

def add_service_table(doc, class_name, base_url, services):
    heading = doc.add_heading(class_name, level=1)

    p = doc.add_paragraph()
    run = p.add_run(f'Base URL: {base_url}')
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Larguras das colunas
    widths = [Cm(4), Cm(6.5), Cm(5.5), Cm(5.5), Cm(5)]
    for i, width in enumerate(widths):
        table.columns[i].width = width

    # Cabecalho
    headers = ['Servico', 'Descricao', 'Parametro de Entrada', 'Parametro de Saida', 'Endereco do Endpoint']
    header_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = HEADER_FG
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, HEADER_BG)

    # Linhas de dados
    for svc in services:
        row = table.add_row()
        cells = row.cells

        data = [
            svc['metodo'],
            svc['descricao'],
            svc['entrada'],
            svc['saida'],
            svc['endpoint']
        ]

        for i, text in enumerate(data):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(8.5)
            if i == 0:
                run.bold = True
                run.font.size = Pt(8.5)
            if i == 4:
                run.font.size = Pt(8)

    doc.add_paragraph('')


# ============================================
# UCEnderecoGeralServicos
# ============================================
endereco_services = [
    {
        'metodo': 'obterEnderecoPorID',
        'descricao': 'Busca um endereco no banco de dados pelo seu ID.',
        'entrada': 'int id - ID do endereco',
        'saida': 'Sucesso: Endereco (JSON)\n404: "Endereco nao encontrado"\nErro: EnderecoException("ID invalido")',
        'endpoint': 'GET /endereco/id/{id}'
    },
    {
        'metodo': 'obterEnderecoPorCEP',
        'descricao': 'Busca um endereco no banco de dados local pelo CEP.',
        'entrada': 'String cep - CEP (8 digitos)',
        'saida': 'Sucesso: Endereco (JSON)\n404: "CEP nao encontrado"\nErro: EnderecoException("Formato de CEP invalido")',
        'endpoint': 'GET /endereco/cep/{cep}'
    },
    {
        'metodo': 'obterCidade',
        'descricao': 'Busca uma cidade pelo seu ID.',
        'entrada': 'int idCidade - ID da cidade',
        'saida': 'Sucesso: Cidade (JSON)\n404: "Cidade nao encontrada"\nErro: EnderecoException("ID de cidade invalido")',
        'endpoint': 'GET /endereco/cidade/{id}'
    },
    {
        'metodo': 'obterEnderecoExterno',
        'descricao': 'Consulta o CEP na API externa ViaCEP e retorna os dados do endereco.',
        'entrada': 'String cep - CEP (8 digitos)',
        'saida': 'Sucesso: Endereco (JSON do ViaCEP)\nErro: EnderecoException("Erro ao consultar API externa")',
        'endpoint': 'GET /endereco/externo/{cep}'
    },
    {
        'metodo': 'cadastrarEndereco',
        'descricao': 'Cadastra um novo endereco no banco de dados.',
        'entrada': 'Endereco (JSON) - com cep, logradouro, bairro, cidade',
        'saida': 'Sucesso: Endereco criado (JSON)\n400: "Dados do endereco invalidos ou incompletos"\n400: "Endereco ja cadastrado"',
        'endpoint': 'POST /endereco/criar'
    },
]

# ============================================
# UCOrdemServicos
# ============================================
ordem_services = [
    {
        'metodo': 'obterClientePorCPF',
        'descricao': 'Busca um cliente pelo CPF.',
        'entrada': 'String cpf - CPF do cliente (11 digitos)',
        'saida': 'Sucesso: Cliente (JSON)\n404: "Cliente nao encontrado"\n400: OrdemServicoException("Formato de CPF invalido")',
        'endpoint': 'GET /cliente/cpf/{cpf}'
    },
    {
        'metodo': 'cadastrarCliente',
        'descricao': 'Cadastra um novo cliente com endereco, telefones e emails.',
        'entrada': 'Cliente (JSON) - com cpf, nome, emails[], telefones[], endereco',
        'saida': 'Sucesso: Cliente criado (JSON)\n400: "Formato de CPF invalido"\n400: "Nome do cliente e obrigatorio"\n400: "Cliente com este CPF ja esta cadastrado"',
        'endpoint': 'POST /cliente/criar'
    },
    {
        'metodo': 'listarAtendentes',
        'descricao': 'Lista todos os atendentes cadastrados no sistema.',
        'entrada': 'Nenhum',
        'saida': 'Sucesso: List<Atendente> (JSON)\n500: OrdemServicoException("Erro ao listar atendentes")',
        'endpoint': 'GET /atendente'
    },
    {
        'metodo': 'obterAtendentePorId',
        'descricao': 'Busca um atendente pelo seu ID.',
        'entrada': 'int id - ID do atendente',
        'saida': 'Sucesso: Atendente (JSON)\n404: "Atendente nao encontrado"\n500: OrdemServicoException("Erro ao buscar atendente")',
        'endpoint': 'GET /atendente/{id}'
    },
    {
        'metodo': 'listarServicos',
        'descricao': 'Lista todos os tipos de servico disponiveis.',
        'entrada': 'Nenhum',
        'saida': 'Sucesso: List<Servico> (JSON)\n500: OrdemServicoException("Erro ao listar servicos")',
        'endpoint': 'GET /servico'
    },
    {
        'metodo': 'listarOrdensServico',
        'descricao': 'Lista todas as ordens de servico.',
        'entrada': 'Nenhum',
        'saida': 'Sucesso: List<OrdemServico> (JSON)\n500: OrdemServicoException("Erro ao listar ordens de servico")',
        'endpoint': 'GET /ordem-servico'
    },
    {
        'metodo': 'obterOrdemServicoPorNumero',
        'descricao': 'Busca uma ordem de servico pelo numero.',
        'entrada': 'int nroOrdem - Numero da OS',
        'saida': 'Sucesso: OrdemServico (JSON)\n404: "Ordem de servico nao encontrada"\n500: OrdemServicoException("Erro ao buscar OS")',
        'endpoint': 'GET /ordem-servico/{nro}'
    },
    {
        'metodo': 'criarOrdemServico',
        'descricao': 'Cria uma nova ordem de servico com servicos associados. Define data de emissao e calcula total automaticamente.',
        'entrada': 'OrdemServico (JSON) - com descricao, cliente, atendente, servicos[]',
        'saida': 'Sucesso: OrdemServico criada (JSON)\n400: "Dados da ordem de servico invalidos"\n500: OrdemServicoException("Erro ao criar OS")',
        'endpoint': 'POST /ordem-servico/criar'
    },
]

# ============================================
# UCReceitaServicos
# ============================================
receita_services = [
    {
        'metodo': 'obterPacientePorCPF',
        'descricao': 'Busca um paciente pelo CPF.',
        'entrada': 'String cpf - CPF do paciente (11 digitos)',
        'saida': 'Sucesso: Paciente (JSON)\n404: "Paciente nao encontrado"\n400: ReceitaException("Formato de CPF invalido")',
        'endpoint': 'GET /paciente/cpf/{cpf}'
    },
    {
        'metodo': 'cadastrarPaciente',
        'descricao': 'Cadastra um novo paciente com endereco, telefones e emails.',
        'entrada': 'Paciente (JSON) - com cpf, nome, emails[], telefones[], endereco',
        'saida': 'Sucesso: Paciente criado (JSON)\n400: "Formato de CPF invalido"\n400: "Nome do paciente e obrigatorio"\n400: "Paciente com este CPF ja esta cadastrado"',
        'endpoint': 'POST /paciente/criar'
    },
    {
        'metodo': 'listarMedicos',
        'descricao': 'Lista todos os medicos cadastrados.',
        'entrada': 'Nenhum',
        'saida': 'Sucesso: List<Medico> (JSON)\n500: ReceitaException("Erro ao listar medicos")',
        'endpoint': 'GET /medico'
    },
    {
        'metodo': 'obterMedicoPorId',
        'descricao': 'Busca um medico pelo seu ID.',
        'entrada': 'int id - ID do medico',
        'saida': 'Sucesso: Medico (JSON)\n404: "Medico nao encontrado"\n500: ReceitaException("Erro ao buscar medico")',
        'endpoint': 'GET /medico/{id}'
    },
    {
        'metodo': 'obterCIDByCodigo',
        'descricao': 'Busca um CID pelo codigo. Tenta match exato e depois LIKE.',
        'entrada': 'String codigo - Codigo CID (ex: J00, G43)',
        'saida': 'Sucesso: CID (JSON)\n404: "CID nao encontrado"\n500: ReceitaException("Erro ao buscar CID")',
        'endpoint': 'GET /cid/{codigo}'
    },
    {
        'metodo': 'listarMedicamentos',
        'descricao': 'Lista medicamentos. Filtra por nome se parametro fornecido, senao lista todos.',
        'entrada': 'String nome (opcional) - filtro por nome do medicamento',
        'saida': 'Sucesso: List<Medicamento> (JSON)\n500: ReceitaException("Erro ao listar medicamentos")',
        'endpoint': 'GET /medicamento\nGET /medicamento?nome={nome}'
    },
    {
        'metodo': 'listarReceitasMedicas',
        'descricao': 'Lista todas as receitas medicas.',
        'entrada': 'Nenhum',
        'saida': 'Sucesso: List<ReceitaMedica> (JSON)\n500: ReceitaException("Erro ao listar receitas medicas")',
        'endpoint': 'GET /receita-medica'
    },
    {
        'metodo': 'obterReceitaMedicaPorNumero',
        'descricao': 'Busca uma receita medica pelo numero.',
        'entrada': 'int numeroReceita - Numero da receita',
        'saida': 'Sucesso: ReceitaMedica (JSON)\n404: "Receita nao encontrada"\n500: ReceitaException("Erro ao buscar receita medica")',
        'endpoint': 'GET /receita-medica/{numero}'
    },
    {
        'metodo': 'criarReceitaMedica',
        'descricao': 'Cria uma nova receita medica com prescricoes. Define data de emissao automaticamente.',
        'entrada': 'ReceitaMedica (JSON) - com paciente, medico, cid, prescricoes[]',
        'saida': 'Sucesso: ReceitaMedica criada (JSON)\n400: "Dados da receita medica invalidos"\n500: ReceitaException("Erro ao criar receita medica")',
        'endpoint': 'POST /receita-medica/criar'
    },
]

# Gerar as tabelas
add_service_table(doc, 'UCEnderecoGeralServicos', 'http://localhost:8080/MyEnderecoServicos/resources', endereco_services)
add_service_table(doc, 'UCOrdemServicos', 'http://localhost:8080/MyOrdemServicos/resources', ordem_services)
add_service_table(doc, 'UCReceitaServicos', 'http://localhost:8080/MyReceitaServicos/resources', receita_services)

output_path = r'C:\Users\Hugo\Desktop\Codes\NetBeans\docs\Tabela_Servicos.docx'
doc.save(output_path)
print(f'Documento gerado com sucesso: {output_path}')
